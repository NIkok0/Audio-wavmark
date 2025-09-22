# Standard library imports
import os

# Third-party imports
from flask import current_app


def embed(input_file, watermark, algorithm):
    """文本水印嵌入 - 负责算法调用和文件保存"""
    
    # 输入验证
    if not watermark or not watermark.strip():
        raise ValueError("水印内容不能为空")
    
    if len(watermark) > 10:
        raise ValueError("水印内容不能超过10个字符")
    
    # 获取文件扩展名
    _, extension = os.path.splitext(input_file)
    extension = extension[1:].lower()
    
    # 生成函数名 (格式: embed_扩展名_算法名)
    function_name = f"embed_{extension}_{algorithm.lower()}"
    
    try:
        # 获取当前模块中的函数
        embed_function = globals().get(function_name)
        if embed_function is None:
            raise ValueError(f"文本水印算法 {algorithm} 不支持 {extension} 格式")
        
        # 调用算法，获取处理后的文本内容
        full_path = embed_function(input_file, watermark)
        
        # 文件保存逻辑
        # original_name = os.path.basename(input_file)
        # name_without_ext = os.path.splitext(original_name)[0]
        # filename = f"{name_without_ext}_embed.{extension}"
        #
        # # 从app.config获取保存路径
        # embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
        # full_path = os.path.join(embed_dir, filename)
        #
        # # 保存文件
        # with open(full_path, 'w', encoding='utf-8') as f:
        #     f.write(processed_text)
        
        return full_path  # 返回完整路径
        
    except Exception as e:
        error_msg = f"文本水印算法 {algorithm} 失败: {str(e)}"
        print(error_msg)
        raise ValueError(error_msg) from e

def extract(input_file, algorithm):
    """文本水印提取 - 支持多种算法（基于配置）"""
    # 获取文件扩展名
    _, extension = os.path.splitext(input_file)
    extension = extension[1:].lower()
    
    # 生成函数名 (格式: extract_扩展名_算法名)
    function_name = f"extract_{extension}_{algorithm.lower()}"
    
    try:
        # 获取当前模块中的函数
        extract_function = globals().get(function_name)
        if extract_function is None:
            raise ValueError(f"文本水印算法 {algorithm} 不支持 {extension} 格式")
        
        return extract_function(input_file)
        
    except Exception as e:
        error_msg = f"文本水印提取算法 {algorithm} 失败: {str(e)}"
        print(error_msg)
        raise ValueError(error_msg) from e

# TXT格式
def embed_txt_zbit(input_file, watermark):
    """TXT格式专用"""

    def watermark_to_zwc(watermark):
        """将字符串水印编码为零宽字符序列（二进制编码）"""
        # 对于中文字符，需要用UTF-8编码而不是直接用ord()
        utf8_bytes = watermark.encode('utf-8')
        binary = ''.join(format(byte, '08b') for byte in utf8_bytes)
        return ''.join('\u200B' if bit == '0' else '\u200D' for bit in binary)

    def embed_watermark(input_path, watermark):
        """在txt中嵌入不可见水印"""
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        zwc = watermark_to_zwc(watermark)

        # 插入位置：文件末尾，也可以改为某些固定模式后
        marked_content = content + '\n' + zwc

        original_name = os.path.basename(input_file)
        name_without_ext = os.path.splitext(original_name)[0]
        filename = f"{name_without_ext}_embed.{'txt'}"

        # 从app.config获取保存路径
        embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
        full_path = os.path.join(embed_dir, filename)

        # 保存文件
        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(marked_content)

        return full_path

    full_path=embed_watermark(input_file, watermark)
    return full_path

def extract_txt_zbit(input_file):
    """TXT格式专用"""

    def zwc_to_watermark(zwc_text):
        """从零宽字符序列中提取水印"""
        binary = ''
        for c in zwc_text:
            if c == '\u200B':
                binary += '0'
            elif c == '\u200D':
                binary += '1'
        
        # 确保二进制长度是8的倍数
        if len(binary) % 8 != 0:
            # 截断到最近的8的倍数
            binary = binary[:-(len(binary) % 8)] if len(binary) >= 8 else ''
        
        if not binary:
            return None
            
        try:
            # 将二进制转换为字节，然后用UTF-8解码
            bytes_data = []
            for i in range(0, len(binary), 8):
                byte_val = int(binary[i:i + 8], 2)
                bytes_data.append(byte_val)
            
            # 用UTF-8解码字节序列
            return bytes(bytes_data).decode('utf-8')
        except (ValueError, UnicodeDecodeError):
            return None

    def extract_watermark(path):
        """从txt中提取水印（零宽字符）"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            raise ValueError(f"无法读取文件 {path}: {str(e)}")

        # 提取所有零宽字符
        zwc_chars = ''.join(c for c in content if c in ('\u200B', '\u200D'))
        
        if not zwc_chars:
            raise ValueError("文件中未找到零宽字符水印")
            
        watermark = zwc_to_watermark(zwc_chars)
        
        if not watermark:
            raise ValueError("零宽字符解码失败，可能不是有效的水印")
            
        return watermark

    print("text_watermark_extract for TXT!")
    return extract_watermark(input_file)

# DOC格式
def embed_doc_space(input_file, watermark):
    """DOC格式专用"""
    import win32com.client
    import os

    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'doc'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    # 保存文件
    def embed_watermark_in_char_spacing(doc_path, output_path, watermark_text):
        """
        将 watermark_text 以二进制形式嵌入到 doc 文档中：
          - 把每个字符的 Font.Spacing 设置为 0 或 1，分别表示二进制 0/1。
          - 如果文档中字符数量不够，则自动在末尾插入空格补齐。
        参数：
          doc_path       - 源 .doc/.docx 文档路径
          output_path    - 嵌入后保存的文档路径
          watermark_text - 要隐藏的字符串（水印内容）
        """
        # 1. 转为绝对路径
        doc_path = os.path.abspath(doc_path)
        output_path = os.path.abspath(output_path)

        # 2. 把 watermark_text 转为二进制序列（每字符 8 位）
        binary_data = ''.join(format(ord(c), '08b') for c in watermark_text)
        total_needed = len(binary_data)  # 需要嵌入的位数

        # 3. 启动 Word 后台打开文档
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)

        # 4. 拿到整个文档的字符集合
        # full_range = doc.Range(0, 0).GoTo(What=win32com.client.constants.wdGoToLine, Which=win32com.client.constants.wdGoToLast)
        # 事实上 doc.Range() 就可以代表全文，只是下面我们直接从 doc.Range().Characters 获取
        chars = doc.Range().Characters
        total_chars = chars.Count

        # 5. 如果字符不足，就在文档尾部插入空格以补齐
        if total_needed > total_chars:
            extra_needed = total_needed - total_chars
            # 在文档末尾插入 extra_needed 个空格，使得 Characters.Count 增加
            doc.Content.InsertAfter(" " * extra_needed)
            # 重新获取 Characters 集合
            chars = doc.Range().Characters
            total_chars = chars.Count

        # 6. 遍历二进制位，把前 total_needed 个字符的 Font.Spacing 设置为 0/1
        #    注意：Characters 集合是 1-based，下标 i+1 对应二进制的第 i 位
        for i, bit in enumerate(binary_data):
            char_range = chars(i + 1)  # 第 i+1 个字符
            if bit == '0':
                char_range.Font.Spacing = 0  # 二进制 0 → 间距 0 pt
            else:
                char_range.Font.Spacing = 1  # 二进制 1 → 间距 1 pt

        # 7. 保存并退出
        doc.SaveAs(output_path)
        doc.Close()
        word.Quit()
        return output_path

    return embed_watermark_in_char_spacing(input_file,full_path, watermark)

def extract_doc_space(input_file):
    """DOC格式专用"""
    import win32com.client
    import pythoncom
    import os
    def extract_watermark_from_char_spacing(doc_path, length=6):
        """
        从之前用 embed_watermark_in_char_spacing 嵌入的数据中提取水印。
        参数：
          doc_path - 含隐藏水印的文档路径
          length   - 原始水印的字符长度（例如 "own" → length=3；“owner”→ length=5）
        返回：
          提取出的原始字符串
        """
        # 1. 转为绝对路径，打开 Word 文档
        doc_path = os.path.abspath(doc_path)
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)

        # 2. 需要读取的二进制总位数 = length * 8
        bits_needed = length * 8
        chars = doc.Range().Characters

        # 3. 如果字符数量小于 bits_needed，说明读取不完整，可以抛错或提示
        if bits_needed > chars.Count:
            doc.Close()
            word.Quit()
            raise ValueError(f"文档字符数 ({chars.Count}) 少于水印读取需要的位数 ({bits_needed})，无法提取完整水印。")

        # 4. 逐位读取字符的 Font.Spacing 属性，Build 二进制串
        binary = []
        for i in range(bits_needed):
            char_range = chars(i + 1)  # 第 i+1 个字符
            spacing = char_range.Font.Spacing
            # 如果 spacing 大于 0（我们嵌入时一律设为 1），则视为 '1'，否则 '0'
            bit = '1' if spacing and spacing > 0 else '0'
            binary.append(bit)
        binary_str = ''.join(binary)

        # 5. 将二进制按照每 8 位拆分并转为字符
        chars_out = []
        for i in range(0, len(binary_str), 8):
            byte = binary_str[i:i + 8]
            chars_out.append(chr(int(byte, 2)))
        watermark = ''.join(chars_out)

        doc.Close()
        word.Quit()
        print(f"🔍 提取出的水印: {watermark}")
        return watermark
    return extract_watermark_from_char_spacing(input_file, 6)


def embed_docx_space(input_file, watermark):
    """DOCX格式专用 - 使用python-docx库"""
    from docx import Document
    import os

    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'docx'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    def watermark_to_zwc(watermark):
        """将字符串水印编码为零宽字符序列（二进制编码）"""
        utf8_bytes = watermark.encode('utf-8')
        binary = ''.join(format(byte, '08b') for byte in utf8_bytes)
        return ''.join('\u200B' if bit == '0' else '\u200D' for bit in binary)

    def embed_watermark_in_docx(doc_path, output_path, watermark_text):
        """在DOCX文档中嵌入零宽字符水印"""
        doc = Document(doc_path)
        zwc = watermark_to_zwc(watermark_text)
        
        # 在第一个段落末尾添加零宽字符水印
        if doc.paragraphs:
            first_paragraph = doc.paragraphs[0]
            if first_paragraph.runs:
                first_paragraph.runs[-1].text += zwc
            else:
                first_paragraph.add_run(zwc)
        else:
            # 如果没有段落，创建一个新段落
            p = doc.add_paragraph()
            p.add_run(zwc)
        
        doc.save(output_path)
        print(f"水印已嵌入DOCX文档: {output_path}")
        return output_path

    return embed_watermark_in_docx(input_file, full_path, watermark)

def extract_docx_space(input_file):
    """DOCX格式专用 - 使用python-docx库提取零宽字符水印"""
    from docx import Document
    import os

    def zwc_to_watermark(zwc_text):
        """从零宽字符序列中提取水印"""
        binary = ''
        for c in zwc_text:
            if c == '\u200B':
                binary += '0'
            elif c == '\u200D':
                binary += '1'
        
        # 确保二进制长度是8的倍数
        if len(binary) % 8 != 0:
            binary = binary[:-(len(binary) % 8)] if len(binary) >= 8 else ''
        
        if not binary:
            return None
            
        try:
            # 将二进制转换为字节，然后用UTF-8解码
            bytes_data = []
            for i in range(0, len(binary), 8):
                byte_val = int(binary[i:i + 8], 2)
                bytes_data.append(byte_val)
            
            return bytes(bytes_data).decode('utf-8')
        except (ValueError, UnicodeDecodeError):
            return None

    def extract_watermark_from_docx(doc_path):
        """从DOCX文档中提取零宽字符水印"""
        try:
            doc = Document(doc_path)
        except Exception as e:
            raise ValueError(f"无法读取DOCX文件 {doc_path}: {str(e)}")

        # 从所有段落和运行中查找零宽字符
        zwc_chars = ''
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    zwc_chars += ''.join(c for c in run.text if c in ('\u200B', '\u200D'))
        
        if not zwc_chars:
            raise ValueError("DOCX文件中未找到零宽字符水印")
            
        watermark = zwc_to_watermark(zwc_chars)
        
        if not watermark:
            raise ValueError("零宽字符解码失败，可能不是有效的水印")
            
        return watermark

    print("text_watermark_extract for DOCX!")
    return extract_watermark_from_docx(input_file)

def embed_xml_space(input_file, watermark):
    """XML格式专用"""
    from collections import Counter
    import xml.etree.ElementTree as ET
    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'xml'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    # 保存文件
    def watermark_to_zwc(watermark):
        binary = ''.join(format(ord(c), '08b') for c in watermark)
        return ''.join('\u200B' if b == '0' else '\u200D' for b in binary)
    def embed_watermark_in_xml(input_file, output_file, watermark, interval=2):
        tree = ET.parse(input_file)
        root = tree.getroot()
        zwc = watermark_to_zwc(watermark)

        count = 0
        for elem in root.iter():
            if elem.text and elem.text.strip():
                count += 1
                if count % interval == 0:
                    elem.text += zwc

        tree.write(output_file, encoding="utf-8", xml_declaration=True)
        print(f"✅ 已将水印嵌入 XML: {output_file}")

    embed_watermark_in_xml(input_file, full_path, watermark)
    return full_path

def extract_xml_space(input_file):
    """XML格式专用"""
    from collections import Counter
    import xml.etree.ElementTree as ET
    def zwc_to_watermark(zwc_text):
        binary = ''.join('0' if c == '\u200B' else '1' for c in zwc_text if c in ('\u200B', '\u200D'))
        
        # 确保二进制长度是8的倍数
        if len(binary) % 8 != 0:
            binary = binary[:-(len(binary) % 8)] if len(binary) >= 8 else ''
        
        if not binary:
            return None
            
        try:
            # 将二进制转换为字节，然后用UTF-8解码
            bytes_data = []
            for i in range(0, len(binary), 8):
                byte_val = int(binary[i:i + 8], 2)
                bytes_data.append(byte_val)
            
            # 用UTF-8解码字节序列
            return bytes(bytes_data).decode('utf-8')
        except (ValueError, UnicodeDecodeError):
            return None

    def extract_watermark_from_xml(path):
        tree = ET.parse(path)
        root = tree.getroot()

        zwc_fragments = []

        for elem in root.iter():
            if elem.text:
                zwc = ''.join(c for c in elem.text if c in ('\u200B', '\u200D'))
                if zwc:
                    zwc_fragments.append(zwc)

        if not zwc_fragments:
            print("未检测到零宽字符水印")
            return None

        decoded = [zwc_to_watermark(fragment) for fragment in zwc_fragments]
        most_common = Counter(decoded).most_common(1)

        if most_common:
            print(f"提取到水印: {most_common[0][0]}")
            return most_common[0][0]
        else:
            print("水印格式无效")
            return None

    print("text_watermark_extract for XML!")
    return extract_watermark_from_xml(input_file)




def embed_pdf_md(input_file, watermark):
    """PDF格式专用"""
    import fitz  # PyMuPDF
    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'pdf'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)
    def embed_invisible_watermark(input_pdf, output_pdf, watermark_text, x=20, y=20):
        doc = fitz.open(input_pdf)
        for page in doc:
            page.insert_text(
                (x, y),
                watermark_text,
                fontsize=0.5,  # 非常小
                color=(1, 1, 1),  # 白色（或透明）
                fill_opacity=0.0,  # 完全透明
                render_mode=3  # 不参与选择/复制
            )
        doc.save(output_pdf)
        doc.close()

    embed_invisible_watermark(input_file, full_path, watermark)
    return full_path

def extract_pdf_md(input_file):
    """PDF格式专用"""
    import fitz  # PyMuPDF
    def extract_watermark_text_by_position(pdf_path, x_range=(15, 25), y_range=(15, 25)):
        doc = fitz.open(pdf_path)
        extracted = []
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        x0, y0 = span["bbox"][0], span["bbox"][1]
                        if x_range[0] <= x0 <= x_range[1] and y_range[0] <= y0 <= y_range[1]:
                            extracted.append(span["text"])
        doc.close()

        if not extracted:
            print("未检测到水印文本")
        else:
            extracted_set = set(extracted)
            watermark = ''.join(extracted_set)
            print(f"提取到的水印: {watermark}")
            return watermark

    print("text_watermark_extract for PDF!")
    return extract_watermark_text_by_position(input_file, x_range=(15, 25), y_range=(15, 25))


def embed_md_zbit(input_file, watermark):
    """MD格式专用"""
    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'md'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    def watermark_to_zwc(watermark):
        """将字符串水印编码为零宽字符序列（二进制编码）"""
        # 对于中文字符，需要用UTF-8编码而不是直接用ord()
        utf8_bytes = watermark.encode('utf-8')
        binary = ''.join(format(byte, '08b') for byte in utf8_bytes)
        return ''.join('\u200B' if bit == '0' else '\u200D' for bit in binary)
    def embed_watermark(input_path, output_path, watermark):
        """在txt中嵌入不可见水印"""
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        zwc = watermark_to_zwc(watermark)

        # 插入位置：文件末尾，也可以改为某些固定模式后
        marked_content = content + '\n' + zwc

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(marked_content)

        print(f"已嵌入水印到: {output_path}")

    embed_watermark(input_file, full_path, watermark)
    return full_path

def extract_md_zbit(input_file):
    """MD格式专用"""

    def zwc_to_watermark(zwc_text):
        """从零宽字符序列中提取水印"""
        binary = ''
        for c in zwc_text:
            if c == '\u200B':
                binary += '0'
            elif c == '\u200D':
                binary += '1'
        
        # 确保二进制长度是8的倍数
        if len(binary) % 8 != 0:
            binary = binary[:-(len(binary) % 8)] if len(binary) >= 8 else ''
        
        if not binary:
            return None
            
        try:
            # 将二进制转换为字节，然后用UTF-8解码
            bytes_data = []
            for i in range(0, len(binary), 8):
                byte_val = int(binary[i:i + 8], 2)
                bytes_data.append(byte_val)
            
            # 用UTF-8解码字节序列
            return bytes(bytes_data).decode('utf-8')
        except (ValueError, UnicodeDecodeError):
            return None
    def extract_watermark(path):
        """从txt中提取水印（零宽字符）"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 提取所有零宽字符
        zwc_chars = ''.join(c for c in content if c in ('\u200B', '\u200D'))

        if not zwc_chars:
            print("未找到零宽字符水印")
            return None

        watermark = zwc_to_watermark(zwc_chars)
        print(f"提取的水印: {watermark}")
        return watermark

    print("text_watermark_extract for MD!")
    return extract_watermark(input_file)


def embed_sql_zbit(input_file, watermark):
    """SQL格式专用"""
    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'md'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    def watermark_to_zwc(watermark):
        binary = ''.join(format(ord(c), '08b') for c in watermark)
        return ''.join('\u200B' if bit == '0' else '\u200D' for bit in binary)

    def embed_watermark(input_path, output_path, watermark):
        zwc = watermark_to_zwc(watermark)
        inserted = False

        with open(input_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # 先找注释行嵌入水印（--单行注释 或 /*多行注释开头）
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith('--') or stripped.startswith('/*'):
                lines[i] = line.rstrip('\n') + zwc + '\n'
                inserted = True
                break

        # 如果没有注释行，找第一条空行插入带水印注释
        if not inserted:
            for i, line in enumerate(lines):
                if line.strip() == '':
                    lines[i] = f'-- {zwc}\n'
                    inserted = True
                    print("未发现注释，已在空行插入注释水印")
                    break

        # 如果依然没插入（没有空行），就在文件头插入一行注释水印
        if not inserted:
            lines.insert(0, f'-- {zwc}\n')
            print("文件无空行，已在文件头插入注释水印")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(lines)

    embed_watermark(input_file, full_path, watermark)
    return full_path

def extract_sql_zbit(input_file):
    """SQL格式专用"""
    import re
    def zwc_to_watermark(zwc_text):
        binary = ''.join('0' if c == '\u200B' else '1' for c in zwc_text if c in ('\u200B', '\u200D'))
        chars = [chr(int(binary[i:i + 8], 2)) for i in range(0, len(binary), 8)]
        return ''.join(chars)

    def extract_watermark(path):
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip().startswith('--') or '/*' in line:
                    zwc_part = ''.join(c for c in line if c in ('\u200B', '\u200D'))
                    if zwc_part:
                        watermark = zwc_to_watermark(zwc_part)
                        print(f"提取到的水印: {watermark}")
                        return watermark
        print("未找到零宽字符水印")
        return None

    print("text_watermark_extract for SQL!")
    return extract_watermark(input_file)



def embed_csv_zbit(input_file, watermark):
    """CSV格式专用"""
    import csv
    from collections import Counter
    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'csv'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    def watermark_to_ws(watermark):
        binary = ''.join(format(ord(c), '08b') for c in watermark)
        return ''.join(' ' if bit == '0' else '\t' for bit in binary)
    def embed_ws_watermark_in_csv(input_path, output_path, watermark, column=2, row_interval=2):
        ws_code = watermark_to_ws(watermark)
        with open(input_path, 'r', encoding='utf-8', newline='') as f:
            rows = list(csv.reader(f))

        for i in range(1, len(rows)):
            if i % row_interval == 0 and column < len(rows[i]):
                rows[i][column] += ws_code  # 直接在字段末尾添加空白字符水印

        with open(output_path, 'w', encoding='utf-8', newline='') as f:
            writer = csv.writer(f)
            writer.writerows(rows)
        print(f"已将水印嵌入至: {output_path}")

    embed_ws_watermark_in_csv(input_file, full_path, watermark)
    return full_path

def extract_csv_zbit(input_file):
    """CSV格式专用"""
    import csv
    from collections import Counter
    def ws_to_watermark(ws_text, length_chars):
        bits = ''.join('0' if c == ' ' else '1' for c in ws_text if c in (' ', '\t'))
        if len(bits) < length_chars * 8:
            return None
        bits = bits[-(length_chars * 8):]  # 只提取尾部的 length_chars 字符
        chars = [chr(int(bits[i:i + 8], 2)) for i in range(0, len(bits), 8)]
        return ''.join(chars)

    def extract_ws_watermark_from_csv(path, watermark_length_chars, column=2):
        ws_segments = []
        with open(path, 'r', encoding='utf-8', newline='') as f:
            for i, row in enumerate(csv.reader(f)):
                if i == 0:
                    continue  # 跳过标题行
                if column < len(row):
                    field = row[column]
                    # 从字段末尾提取连续的空白字符
                    tail_ws = ''
                    for c in reversed(field):
                        if c in (' ', '\t'):
                            tail_ws = c + tail_ws  # 保持正确顺序
                        else:
                            break  # 遇到非空白字符就停止
                    if tail_ws:
                        ws_segments.append(tail_ws)

        decoded = [ws_to_watermark(ws, watermark_length_chars) for ws in ws_segments if ws]
        decoded = [d for d in decoded if d is not None]
        if not decoded:
            print("未能成功提取水印")
            return None

        most_common = Counter(decoded).most_common(1)
        print(f"提取到的水印: {most_common[0][0]}")
        return most_common[0][0]

    print("text_watermark_extract for CSV!")
    # 尝试不同的水印长度，从1-10个字符
    for length in range(1, 11):
        try:
            result = extract_ws_watermark_from_csv(input_file, watermark_length_chars=length)
            if result and result.strip():  # 如果提取到有效内容
                return result
        except Exception:
            continue
    return extract_ws_watermark_from_csv(input_file, watermark_length_chars=7)  # 默认回退








def embed_xlsx_ns(input_file, watermark):
    """XLSX格式专用"""
    from openpyxl import load_workbook
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter
    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'xlsx'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    def add_hidden_watermark_cell(xlsx_path, output_path, watermark_text):
        wb = load_workbook(xlsx_path)
        ws = wb.active  # 你可以改成指定工作表名：wb['Sheet1']

        # 选择一个偏远的单元格（例如最后一行最后一列）
        row = 1000
        col = 200  # 大约是列 "GR"
        cell = ws.cell(row=row, column=col)
        cell.value = watermark_text

        # 设置字体颜色与背景相同，"伪装" 不可见
        # cell.font = Font(color=cell.fill.fgColor.rgb or "FFFFFF")  # 默认白色
        # 可选：你也可以直接指定白色或其他背景色
        cell.font = Font(color="FFFFFF")

        # 隐藏该列和行
        ws.row_dimensions[row].hidden = True
        ws.column_dimensions[get_column_letter(col)].hidden = True

        # 保存到新文件
        wb.save(output_path)
        print(f"水印已写入隐藏单元格：{get_column_letter(col)}{row}，文件输出为：{output_path}")

    add_hidden_watermark_cell(input_file, full_path, watermark)
    return full_path

def extract_xlsx_ns(input_file):
    """XLSX格式专用"""
    import csv
    from collections import Counter
    def extract_hidden_watermark_cell(xlsx_path, row=1000, col=200):
        from openpyxl import load_workbook

        wb = load_workbook(xlsx_path, data_only=True)
        ws = wb.active

        cell = ws.cell(row=row, column=col)
        print(f"提取到隐藏水印：{cell.value}")
        return cell.value

    print("text_watermark_extract for XLSX!")
    return extract_hidden_watermark_cell(input_file)


def embed_xls_ns(input_file, watermark):
    """XLS格式专用"""
    import xlrd
    import xlwt
    from xlutils.copy import copy
    from xlwt import XFStyle, Font
    original_name = os.path.basename(input_file)
    name_without_ext = os.path.splitext(original_name)[0]
    filename = f"{name_without_ext}_embed.{'xls'}"

    # 从app.config获取保存路径
    embed_dir = current_app.config['MEDIA_FOLDERS']['text']['embed']
    full_path = os.path.join(embed_dir, filename)

    def cellname_to_rowcol(cellname):
        """
        将单元格名称（如 'Z100'）转换为行和列索引（0-based）。
        例如：'Z100' -> (99, 25)（行 99，列 25，Z 是第 26 列）。
        """
        if not cellname:
            raise ValueError("单元格名称不能为空")

        # 分离字母和数字部分
        import re
        match = re.match(r'([A-Z]+)(\d+)', cellname, re.I)
        if not match:
            raise ValueError(f"无效的单元格名称: {cellname}")

        col_str, row_str = match.groups()
        # 转换列字母到索引（A=0, B=1, ..., Z=25, AA=26, ...）
        col = 0
        for char in col_str.upper():
            col = col * 26 + (ord(char) - ord('A') + 1)
        col -= 1  # 转换为 0-based 索引
        row = int(row_str) - 1  # 转换为 0-based 索引

        return row, col

    def embed_hidden_watermark(input_xls, output_xls, watermark_text, sheet_name="Sheet1", cell="Z100"):
        # 打开现有的 XLS 文件
        try:
            rb = xlrd.open_workbook(input_xls, formatting_info=True)
        except Exception as e:
            print(f"❌ 无法打开输入 XLS 文件: {e}")
            return

        rs = rb.sheet_by_name(sheet_name)

        # 复制原始工作簿
        wb = copy(rb)
        ws = wb.get_sheet(sheet_name)

        # 创建新的字体样式，将字体颜色设置为白色（与背景色相同）
        font = Font()
        font.colour_index = 0x9  # 白色
        style = XFStyle()
        style.font = font

        # 写入水印文本
        try:
            row, col = cellname_to_rowcol(cell)
            ws.write(row, col, watermark_text, style)

            # 隐藏所在行和列
            ws.row(row).hidden = True
            ws.col(col).hidden = True
            print(f"水印已嵌入到 {sheet_name} 的单元格 {cell}")
        except Exception as e:
            print(f"嵌入水印失败: {e}")
            return

        # 保存新的 XLS 文件
        try:
            wb.save(output_xls)
            print(f"水印已保存至: {output_xls}")
        except Exception as e:
            print(f"保存 XLS 文件失败: {e}")

    embed_hidden_watermark(input_file, full_path, watermark)
    return full_path

def extract_xls_ns(input_file):
    """XLS格式专用"""
    import csv
    from collections import Counter
    def cellname_to_rowcol(cellname):
        """
        将单元格名称（如 'Z100'）转换为行和列索引（0-based）。
        例如：'Z100' -> (99, 25)（行 99，列 25，Z 是第 26 列）。
        """
        if not cellname:
            raise ValueError("单元格名称不能为空")

        # 分离字母和数字部分
        import re
        match = re.match(r'([A-Z]+)(\d+)', cellname, re.I)
        if not match:
            raise ValueError(f"无效的单元格名称: {cellname}")

        col_str, row_str = match.groups()
        # 转换列字母到索引（A=0, B=1, ..., Z=25, AA=26, ...）
        col = 0
        for char in col_str.upper():
            col = col * 26 + (ord(char) - ord('A') + 1)
        col -= 1  # 转换为 0-based 索引
        row = int(row_str) - 1  # 转换为 0-based 索引

        return row, col
    def extract_hidden_watermark(xls_path, sheet_name="Sheet1", cell="Z100"):
        import xlrd
        # 打开 XLS 文件
        try:
            wb = xlrd.open_workbook(xls_path)
        except Exception as e:
            print(f"❌ 无法打开 XLS 文件: {e}")
            return None

        try:
            ws = wb.sheet_by_name(sheet_name)
        except Exception as e:
            print(f"无法找到工作表 {sheet_name}: {e}")
            return None

        # 获取水印文本
        try:
            row, col = cellname_to_rowcol(cell)
            watermark = ws.cell_value(row, col)
            if watermark:
                print(f"从 {sheet_name} 的单元格 {cell} 提取到水印: {watermark}")
                return watermark
            else:
                print(f"在 {sheet_name} 的单元格 {cell} 未检测到水印")
                return None
        except Exception as e:
            print(f"提取水印失败: {e}")
            return None

    print("text_watermark_extract for XLS!")
    return extract_hidden_watermark(input_file)
