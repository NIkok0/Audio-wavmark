"""
生成 Word 说明（与 README.md 内容对齐）。
在 backend-java 目录执行:  python scripts/gen_readme_docx.py
输出: Watermark-Java-API-Guide.docx（与 README 同步；纯 ASCII 文件名便于各环境）
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(11)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    r = p.add_run(text.strip("\n"))
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
            for p in table.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def add_note(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r0 = p.add_run(title + " ")
    r0.bold = True
    r0.font.size = Pt(11)
    r0.font.color.rgb = RGBColor(0x00, 0x66, 0x99)
    r0.font.name = "Calibri"
    r0._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r1 = p.add_run(body)
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def build() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Java API 后端（backend-java）")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run("多模块 Spring Boot · Session + Redis · OpenAPI · Flyway")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()

    add_para(
        doc,
        "本说明与仓库 backend-java/README.md 内容对应；日常协作建议以 Markdown 为准。"
        " 需要 Word 时在 backend-java 目录执行：python scripts/gen_readme_docx.py"
        " 可重新生成 Watermark-Java-API-Guide.docx。",
    )

    add_heading(doc, "一、速览", level=1)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["工程形态", "Maven 多模块；可执行入口在子模块 web"],
            ["设计文档", "docs/watermark-java-backend-tech-selection.md（技术选型）；docs/DEPLOY-SERVER.md（生产部署）"],
            ["默认端口", "8080"],
            ["OpenAPI", "http://localhost:8080/v3/api-docs"],
            ["Swagger", "http://localhost:8080/swagger-ui.html"],
            ["水印算法", "仓库 watermark/（Python）；异步任务见附录 Worker"],
        ],
    )
    add_para(
        doc,
        "技术栈：Spring Boot 3.2、Java 17、Spring Security 6（BCrypt）、JPA + Flyway + MySQL、"
        "Spring Session Data Redis、springdoc-openapi、AWS SDK v2 S3（MinIO/COS）、腾讯云 STS。",
    )

    add_heading(doc, "最短启动命令", level=2)
    add_para(doc, "在 backend-java 目录执行：")
    add_code(doc, "docker compose up -d\nmvn -pl web -am spring-boot:run")
    add_para(doc, "无 Docker 时跳过第一行，自行保证 MySQL、Redis、MinIO 与 application-dev.yml 或环境变量一致。")
    add_para(doc, "依赖关系示意：MySQL、Redis、MinIO 均连向 Spring Boot API。")

    add_heading(doc, "二、工作目录", level=1)
    add_table(
        doc,
        ["场景", "路径示例"],
        [
            ["相对仓库根", "watermarking/backend-java"],
            ["Windows 示例", "E:\\code\\watermarking\\backend-java"],
        ],
    )
    add_code(doc, "Set-Location E:\\code\\watermarking\\backend-java")
    add_note(
        doc,
        "提示：",
        "若不在该目录，易出现 Could not find the selected project in the reactor: web。"
        " 任意目录可用：mvn -f E:\\code\\watermarking\\backend-java\\pom.xml -pl web -am …",
    )

    add_heading(doc, "三、环境要求", level=1)
    add_table(
        doc,
        ["项", "要求"],
        [
            ["JDK", "17（完整 JDK，JAVA_HOME 指向 JDK）"],
            ["Maven", "3.9+"],
            ["Docker", "推荐：docker compose 起依赖；未安装则见「无 Docker」"],
        ],
    )

    add_heading(doc, "四、第一次启动", level=1)

    add_heading(doc, "步骤 A：启动依赖", level=2)
    add_code(doc, "docker compose up -d")
    add_table(
        doc,
        ["服务", "端口", "说明"],
        [
            ["MySQL", "3306", "库 watermark；用户 root；密码 devpass"],
            ["Redis", "6379", "无密码"],
            ["MinIO API", "9000", "minioadmin / minioadmin；桶 watermark"],
            ["MinIO 控制台", "9001", "Web 管理"],
        ],
    )
    add_para(doc, "首次启动后等待 MySQL 就绪（约数十秒）再启动 Java。")

    add_heading(doc, "无 Docker", level=3)
    add_para(
        doc,
        "若 PowerShell 提示无法将「docker」项识别为名称：未安装 Docker 或未加入 PATH。"
        " 可安装 Docker Desktop for Windows 并重启后重试；或本机自行安装 MySQL 8、Redis、MinIO，"
        "与 application-dev.yml 或 WM_DATASOURCE_*、WM_REDIS_*、WM_MINIO_* 对齐后直接进入步骤 C。",
    )

    add_heading(doc, "步骤 B：环境变量（可选）", level=2)
    add_para(doc, "不配也可运行；application-dev.yml 已与 Compose 默认对齐。")
    add_code(
        doc,
        '$env:WM_DATASOURCE_URL = "jdbc:mysql://localhost:3306/watermark?..."\n'
        '$env:WM_DATASOURCE_USERNAME = "root"\n'
        '$env:WM_DATASOURCE_PASSWORD = "devpass"\n'
        '$env:WM_REDIS_HOST = "localhost"\n'
        '$env:WM_REDIS_PORT = "6379"\n'
        '$env:WM_PROFILE = "dev"',
    )
    add_para(doc, "更多见 web/src/main/resources/application.yml。")

    add_heading(doc, "步骤 C：启动 API（务必带 -pl web）", level=2)
    add_code(doc, "mvn -pl web -am spring-boot:run")
    add_table(
        doc,
        ["参数", "含义"],
        [
            ["-pl web", "仅对 web 执行 spring-boot:run（main 在此）"],
            ["-am", "同时构建 web 依赖的 application、infrastructure、domain"],
        ],
    )
    add_note(
        doc,
        "注意：",
        "勿在根目录执行裸命令 mvn spring-boot:run（父 POM 无 main，报 Unable to find a suitable main class）。"
        " 父 pom 已对 spring-boot-maven-plugin 设 skip=true，web 模块 skip=false。",
    )

    add_heading(doc, "步骤 D：验证", level=2)
    add_table(
        doc,
        ["用途", "URL"],
        [
            ["Swagger", "http://localhost:8080/swagger-ui.html"],
            ["OpenAPI JSON", "http://localhost:8080/v3/api-docs"],
            ["健康检查", "http://localhost:8080/actuator/health"],
            ["Prometheus", "http://localhost:8080/actuator/prometheus"],
        ],
    )

    add_heading(doc, "五、日常命令", level=1)
    add_code(doc, "docker compose down\nmvn -pl web -am package -DskipTests")
    add_para(doc, "可执行 Jar 示例：web/target/web-0.1.0-SNAPSHOT.jar（版本以 pom.xml 为准）。")

    add_heading(doc, "六、测试 API", level=1)
    add_heading(doc, "Swagger（推荐）", level=2)
    add_para(
        doc,
        "1）确认 API 已启动；2）打开 Swagger；3）依次 POST /api/v1/auth/register 与 POST /api/v1/auth/login，"
        "浏览器保存 WMSESSIONID；4）再测文件、STS、files/complete、任务等。",
    )
    add_para(
        doc,
        "多数业务 POST 已豁免 CSRF；POST /api/v1/auth/logout 仍需 X-XSRF-TOKEN 与 Cookie XSRF-TOKEN。",
    )

    add_heading(doc, "curl（bash / Git Bash）", level=2)
    add_code(
        doc,
        "BASE=http://localhost:8080\n"
        "curl -s -c cookies.txt -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"username\":\"demo\",\"email\":\"demo@example.com\",\"password\":\"secret12\"}' \\\n"
        "  \"$BASE/api/v1/auth/register\"\n"
        "curl -s -c cookies.txt -b cookies.txt -H \"Content-Type: application/json\" \\\n"
        "  -d '{\"usernameOrEmail\":\"demo\",\"password\":\"secret12\"}' \\\n"
        "  \"$BASE/api/v1/auth/login\"\n"
        'curl -s -b cookies.txt "$BASE/api/v1/files?page=0&size=10"',
    )

    add_heading(doc, "PowerShell", level=2)
    add_para(
        doc,
        "请使用 curl.exe，避免 curl 被 Invoke-WebRequest 别名占用。"
        " 单行注册示例：curl.exe -s -c cookies.txt -H \"Content-Type: application/json\" "
        '-d "{\\"username\\":\\"demo\\",\\"email\\":\\"demo@example.com\\",\\"password\\":\\"secret12\\"}" '
        "http://localhost:8080/api/v1/auth/register",
    )

    add_heading(doc, "Maven 自动化测试", level=2)
    add_code(doc, "mvn -pl web -am test\nmvn -pl web -am verify")
    add_table(
        doc,
        ["类型", "说明"],
        [
            ["单元测试", "不依赖 Docker"],
            ["EndToEndWatermarkFlowTest", "Testcontainers；本机无 Docker 时跳过；CI 跑全量"],
        ],
    )
    add_para(doc, "CI：.github/workflows/backend-java-ci.yml")

    add_heading(doc, "七、常见问题", level=1)
    add_table(
        doc,
        ["现象", "处理"],
        [
            ["Could not find the selected project in the reactor: web", "cd 到 backend-java 或使用 -f pom.xml"],
            ["Unable to find a suitable main class", "使用 mvn -pl web -am spring-boot:run"],
            ["数据库连接失败", "检查 MySQL 与 WM_DATASOURCE_*"],
            ["Redis 连接失败", "检查 Redis 与 WM_REDIS_*"],
            ["对象存储失败", "检查 MinIO；WM_MINIO_ENDPOINT 默认 http://127.0.0.1:9000"],
        ],
    )

    add_heading(doc, "八、附录", level=1)
    add_heading(doc, "Python 异步 Worker", level=2)
    add_code(doc, "python -m watermark.worker.redis_stream_worker")
    add_para(
        doc,
        "在含 watermark 包的仓库根执行；与 Java 共用 SQLALCHEMY_DATABASE_URI、WM_REDIS_*、WM_JOBS_*、WM_MINIO_* 或 COS。"
        " 详见 watermark/worker/redis_stream_worker.py。",
    )

    add_heading(doc, "与既有 Python / 旧版数据共用 MySQL", level=2)
    add_para(
        doc,
        "Java 使用 BCrypt 与 LegacyAwarePasswordEncoder（Werkzeug 旧哈希可在登录后升级）。共用库时注意 Flyway 顺序与备份。",
    )

    add_heading(doc, "Maven 模块", level=2)
    add_para(doc, "依赖链：web → application → infrastructure → domain。")
    add_table(
        doc,
        ["模块", "职责"],
        [
            ["domain", "JPA 实体"],
            ["infrastructure", "Flyway、Repository、S3/MinIO/COS STS、Redis 任务等"],
            ["application", "认证、文件、任务、管理等"],
            ["web", "WatermarkApplication、REST、Prometheus"],
        ],
    )

    add_heading(doc, "运维与后续决策", level=2)
    add_para(
        doc,
        "Session 与 JWT 长期方案；生产环境对 /actuator/prometheus 的访问控制；"
        "可选 OWASP dependency-check、Snyk、Jib/Buildpacks 镜像。",
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("— 文档结束 —")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "Watermark-Java-API-Guide.docx"
    doc = build()
    doc.save(out)
    print(f"Written: {out}")


if __name__ == "__main__":
    main()
